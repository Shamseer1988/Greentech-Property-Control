"""Builds the bilingual (.docx) rental-agreement document from a
generated agreement's clause list.

python-docx over reportlab: Word's own rendering engine shapes Arabic
natively, so a python-docx-built file displays correctly with zero
reshaping work, unlike a from-scratch reportlab/PDF pipeline which
would need `arabic-reshaper` + `python-bidi` + an embedded Arabic font
purely for this one feature.

The whole document body is one borderless 2-column table — English
left (LTR), Arabic right (RTL, right-aligned). RTL paragraph direction
has no first-class python-docx property; it's set via the documented
oxml workaround (`<w:bidi/>` on the paragraph's `pPr`).
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .attachments import company_logo_path

ARABIC_FONT = "Traditional Arabic"
LATIN_FONT = "Calibri"


def _set_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def _set_run_font(run, *, arabic: bool, size: int = 11, bold: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = ARABIC_FONT if arabic else LATIN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    if arabic:
        # w:cs (complex-script) is what Word actually consults for Arabic
        # glyphs — without it, the Latin default font silently wins and
        # Arabic renders in the wrong typeface (or as tofu, in some Word
        # builds).
        rFonts.set(qn("w:cs"), ARABIC_FONT)
        lang = OxmlElement("w:lang")
        lang.set(qn("w:bidi"), "ar-QA")
        rPr.append(lang)
    else:
        rFonts.set(qn("w:ascii"), LATIN_FONT)
        rFonts.set(qn("w:hAnsi"), LATIN_FONT)


def _fill_cell(cell, text: str, *, arabic: bool, bold: bool = False, size: int = 11) -> None:
    """Write `text` into a table cell, one paragraph per blank-line-
    separated block, in the given language/direction."""
    lines = (text or "").split("\n\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if arabic else WD_ALIGN_PARAGRAPH.LEFT
        if arabic:
            _set_rtl(p)
        run = p.add_run(line)
        _set_run_font(run, arabic=arabic, size=size, bold=bold)


def _strip_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _add_row(table, en_text: str, ar_text: str, *, en_bold=False, ar_bold=False,
            en_heading: str | None = None, ar_heading: str | None = None,
            size: int = 11) -> None:
    row = table.add_row()
    left, right = row.cells[0], row.cells[1]
    if en_heading:
        _fill_cell(left, en_heading, arabic=False, bold=True, size=size + 1)
        p = left.add_paragraph()
        run = p.add_run(en_text)
        _set_run_font(run, arabic=False, size=size, bold=en_bold)
    else:
        _fill_cell(left, en_text, arabic=False, bold=en_bold, size=size)
    if ar_heading:
        _fill_cell(right, ar_heading, arabic=True, bold=True, size=size + 1)
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl(p)
        run = p.add_run(ar_text)
        _set_run_font(run, arabic=True, size=size, bold=ar_bold)
    else:
        _fill_cell(right, ar_text, arabic=True, bold=ar_bold, size=size)


def build_docx(*, title_en: str, title_ar: str, agreement_number: str,
               context: dict, clauses: list[dict]) -> bytes:
    """`context` carries the document-chrome fields (lessor/tenant,
    dates, agreement number) that surround the clause list but aren't
    themselves template-specific clause content — the recital and
    party-intro blocks in particular. `clauses` is exactly what
    `agreement_templates.build_clauses()` returned; the preview
    endpoint and this builder always consume the identical list, so
    they can never show different text."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(1.8)

    logo_path = company_logo_path()
    if logo_path:
        try:
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_p.add_run().add_picture(logo_path, width=Cm(2.5))
        except Exception:  # noqa: BLE001 — cosmetic only, must never block generation
            pass

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_en = title_p.add_run(title_en)
    _set_run_font(run_en, arabic=False, size=15, bold=True)
    title_p.add_run("   —   ")
    run_ar = title_p.add_run(title_ar)
    _set_run_font(run_ar, arabic=True, size=15, bold=True)

    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_p.add_run(f"{agreement_number}")
    _set_run_font(ref_run, arabic=False, size=9)

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _strip_table_borders(table)

    lessor, tenant = context["lessor"], context["tenant"]

    # Party-intro blocks — one row per party, English left / Arabic right.
    for party in (lessor, tenant):
        intro_en = (
            f"{party['role_label_en']}: {party['name']}"
            + (f" — CR/registration no. {party['cr_number']}" if party.get("cr_number") else "")
            + (f"\nRepresented by: {party['signatory_name']}"
               f" — ID card no. {party['signatory_id_number']}"
               if party.get("signatory_name") else "")
            + (f", {party['signatory_title']}" if party.get("signatory_title") else "")
            + (f"\nMobile: {party['signatory_mobile'] or party.get('mobile') or ''}"
               if (party.get("signatory_mobile") or party.get("mobile")) else "")
        )
        intro_ar = (
            f"{party['role_label_ar']}: {party['name_ar'] or party['name']}"
            + (f" — قيد رقم {party['cr_number']}" if party.get("cr_number") else "")
            + (f"\nيمثلها: {party['signatory_name_ar'] or party['signatory_name']}"
               f" — بطاقة شخصية رقم {party['signatory_id_number']}"
               if party.get("signatory_name") else "")
            + (f" ({party['signatory_title']})" if party.get("signatory_title") else "")
            + (f"\nجوال: {party['signatory_mobile'] or party.get('mobile') or ''}"
               if (party.get("signatory_mobile") or party.get("mobile")) else "")
        )
        _add_row(table, intro_en, intro_ar)

    recital_en = (
        f"On this {context.get('weekday_en', '')} {context.get('today_str', '')}, "
        f"it is agreed between the above two parties, who — having acknowledged their "
        f"legal capacity to contract and dispose — agreed as follows:"
    )
    recital_ar = (
        f"إنه في يوم {context.get('weekday_ar', '')} الموافق {context.get('today_str', '')} "
        f"قد تم الاتفاق بين الطرفين المذكورين أعلاه، وبعد أن أقر الطرفان بأهليتهما للتعاقد "
        f"والتصرف، اتفقا على ما يلي:"
    )
    _add_row(table, recital_en, recital_ar, en_bold=True, ar_bold=True)

    for clause in clauses:
        _add_row(
            table, clause["body_en"], clause["body_ar"],
            en_heading=clause.get("heading_en"), ar_heading=clause.get("heading_ar"),
        )

    sig_en = (
        f"Party 1 ({lessor['role_label_en']})\n{lessor['name']}\n\n"
        f"Party 2 ({tenant['role_label_en']})\n{tenant['name']}"
    )
    sig_ar = (
        f"الطرف الأول ({lessor['role_label_ar']})\n{lessor['name_ar'] or lessor['name']}\n\n"
        f"الطرف الثاني ({tenant['role_label_ar']})\n{tenant['name_ar'] or tenant['name']}"
    )
    _add_row(table, sig_en, sig_ar, en_bold=True, ar_bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
